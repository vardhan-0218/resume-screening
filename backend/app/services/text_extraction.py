import io
import PyPDF2
import pdfplumber
from docx import Document
import aiofiles
from typing import Union
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

class TextExtractionService:
    """Service for extracting text from various file formats."""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.doc', '.docx', '.txt']
    
    async def extract_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from uploaded file based on its format."""
        file_extension = Path(filename).suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        try:
            if file_extension == '.pdf':
                return await self._extract_from_pdf(file_content)
            elif file_extension in ['.doc', '.docx']:
                return await self._extract_from_docx(file_content)
            elif file_extension == '.txt':
                return await self._extract_from_txt(file_content)
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            raise ValueError(f"Failed to extract text from {filename}: {str(e)}")
    
    async def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file using multiple methods for better accuracy."""
        text_content = ""
        
        try:
            # Method 1: Using pdfplumber (better for complex layouts)
            with io.BytesIO(file_content) as pdf_file:
                with pdfplumber.open(pdf_file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n"
            
            # If pdfplumber fails, fallback to PyPDF2
            if not text_content.strip():
                with io.BytesIO(file_content) as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    for page in pdf_reader.pages:
                        text_content += page.extract_text() + "\n"
        
        except Exception as e:
            logger.warning(f"PDF extraction error: {str(e)}")
            # Last resort: try PyPDF2 if pdfplumber fails
            try:
                with io.BytesIO(file_content) as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    for page in pdf_reader.pages:
                        text_content += page.extract_text() + "\n"
            except Exception as fallback_error:
                raise ValueError(f"Could not extract text from PDF: {str(fallback_error)}")
        
        return self._clean_text(text_content)
    
    async def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            with io.BytesIO(file_content) as docx_file:
                doc = Document(docx_file)
                text_content = []
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_content.append(cell.text)
                
                return self._clean_text('\n'.join(text_content))
        
        except Exception as e:
            raise ValueError(f"Could not extract text from DOCX: {str(e)}")
    
    async def _extract_from_txt(self, file_content: bytes) -> str:
        """Extract text from plain text file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    text_content = file_content.decode(encoding)
                    return self._clean_text(text_content)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            text_content = file_content.decode('utf-8', errors='ignore')
            return self._clean_text(text_content)
        
        except Exception as e:
            raise ValueError(f"Could not extract text from TXT: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        
        # Remove excessive whitespace and normalize line breaks
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        cleaned_text = '\n'.join(lines)
        
        # Remove excessive spaces
        import re
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
        cleaned_text = re.sub(r'\n\s*\n', '\n\n', cleaned_text)
        
        return cleaned_text.strip()
    
    def get_file_info(self, filename: str, file_size: int) -> dict:
        """Get information about the uploaded file."""
        return {
            "filename": filename,
            "extension": Path(filename).suffix.lower(),
            "size_bytes": file_size,
            "supported": Path(filename).suffix.lower() in self.supported_formats
        }